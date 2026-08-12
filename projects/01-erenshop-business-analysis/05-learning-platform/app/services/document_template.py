"""
document_template.py — Analiz Teknik Dokümanı şablonu

Bu modül "single source of truth" olarak çalışır:
  - Şablon Word dosyasını üretir (Eren indirir, doldurur)
  - Aynı bölüm tanımları evaluation servisi tarafından okunur

Bölüm yapısı sektörde yaygın BRD (Business Requirements Document) +
SRS (Software Requirements Specification) hibrit yapısına dayanır.
"""
from io import BytesIO
from datetime import datetime

# python-docx
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH


# ============================================================
# ŞABLON BÖLÜMLERİ (Single Source of Truth)
#   - key: kısa anahtar (evaluation'da arama yapılır)
#   - title: Word'de görünen başlık
#   - guidance: kullanıcıya yol gösterici açıklama (Word'e italik yazılır)
#   - required: bu bölüm zorunlu mu (boşsa puandan düşer)
#   - min_words: bu bölüm için beklenen minimum kelime sayısı
#   - placeholder: kullanıcının silip yerine yazması beklenen örnek metin
# ============================================================
TEMPLATE_SECTIONS = [
    {
        "key": "dokuman_bilgileri",
        "title": "1. Doküman Bilgileri",
        "guidance": "Doküman versiyonu, hazırlayan kişi, tarih, gözden geçiren bilgileri.",
        "required": True,
        "min_words": 20,
        "placeholder": (
            "Doküman Adı: [Talep Adı] Analiz Teknik Dokümanı\n"
            "Versiyon: 1.0\n"
            "Hazırlayan: [Adınız Soyadınız]\n"
            "Tarih: [GG.AA.YYYY]\n"
            "Gözden Geçiren: [Yönetici Adı]\n"
            "Durum: Taslak / İncelemede / Onaylı"
        ),
    },
    {
        "key": "yonetici_ozeti",
        "title": "2. Yönetici Özeti",
        "guidance": (
            "Talebin amacını, hangi ihtiyacı karşılayacağını ve beklenen iş "
            "değerini 4-6 cümle ile özetleyin. Teknik detaya girmeyin."
        ),
        "required": True,
        "min_words": 60,
        "placeholder": (
            "Bu doküman ... özelliğinin ErenShop platformuna eklenmesini açıklamaktadır. "
            "Mevcut durumda ... eksikliği nedeniyle ... yaşanmaktadır. Önerilen çözüm "
            "ile ... iş değeri elde edilecektir."
        ),
    },
    {
        "key": "amac_kapsam",
        "title": "3. Amaç ve Kapsam",
        "guidance": (
            "3.1 Amaç → bu geliştirme neden yapılıyor?\n"
            "3.2 Kapsam Dahilinde → neler yapılacak?\n"
            "3.3 Kapsam Dışı → neler bu projeye dahil DEĞİL?\n"
            "Kapsam dışı bölümü çok önemlidir, scope creep'i engeller."
        ),
        "required": True,
        "min_words": 80,
        "placeholder": (
            "3.1 Amaç:\nBu geliştirmenin amacı ...\n\n"
            "3.2 Kapsam Dahilinde:\n- ...\n- ...\n\n"
            "3.3 Kapsam Dışı:\n- ...\n- ..."
        ),
    },
    {
        "key": "mevcut_durum",
        "title": "4. Mevcut Durum Analizi",
        "guidance": (
            "Şu anda ErenShop'ta ilgili akış nasıl işliyor? Hangi modüller etkilenecek? "
            "Mevcut tablolardan hangilerine dokunulacak veya yeni tablo gerekecek mi?"
        ),
        "required": True,
        "min_words": 60,
        "placeholder": "Mevcut durumda ...",
    },
    {
        "key": "paydaslar",
        "title": "5. Paydaşlar",
        "guidance": (
            "İş tarafı (sponsor, müşteri rolleri, operasyon ekipleri) ve teknik taraf "
            "(geliştirici, test, devops) ayrı ayrı listelenir."
        ),
        "required": True,
        "min_words": 30,
        "placeholder": (
            "- Ürün Sahibi (Product Owner): ...\n"
            "- Son Kullanıcı: müşteri / admin / pazarlama ekibi\n"
            "- Geliştirici Ekip: backend, frontend\n"
            "- Test Ekibi: ...\n"
            "- Operasyon: ..."
        ),
    },
    {
        "key": "fonksiyonel_gereksinimler",
        "title": "6. Fonksiyonel Gereksinimler",
        "guidance": (
            "Numaralandırılmış olarak yazın. Her gereksinim 'Sistem ... yapabilmelidir' "
            "şeklinde net bir cümle olmalı. FG-01, FG-02 gibi kodlayın."
        ),
        "required": True,
        "min_words": 100,
        "placeholder": (
            "FG-01: Sistem ...\n"
            "FG-02: Sistem ...\n"
            "FG-03: Sistem ...\n"
            "FG-04: Sistem ...\n"
            "FG-05: Sistem ..."
        ),
    },
    {
        "key": "fonksiyonel_olmayan_gereksinimler",
        "title": "7. Fonksiyonel Olmayan Gereksinimler",
        "guidance": (
            "Performans, güvenlik, kullanılabilirlik, ölçeklenebilirlik vs. "
            "Mümkünse ölçülebilir metriklerle (örn: 'sayfa < 2 sn açılmalı')."
        ),
        "required": True,
        "min_words": 50,
        "placeholder": (
            "Performans: ...\n"
            "Güvenlik: ...\n"
            "Kullanılabilirlik: ...\n"
            "Loglama: ..."
        ),
    },
    {
        "key": "is_kurallari",
        "title": "8. İş Kuralları",
        "guidance": (
            "Sistemin uyması gereken iş mantığı kuralları. IK-01, IK-02 şeklinde kodlayın. "
            "Validation kuralları, limitler, koşullar burada yer alır."
        ),
        "required": True,
        "min_words": 80,
        "placeholder": (
            "IK-01: ...\n"
            "IK-02: ...\n"
            "IK-03: ...\n"
            "IK-04: ..."
        ),
    },
    {
        "key": "surec_akisi",
        "title": "9. Süreç Akışı",
        "guidance": (
            "Adım adım kullanıcı/sistem akışı. Mümkünse ana akış (happy path) ve "
            "alternatif/hata akışlarını ayrı yazın."
        ),
        "required": True,
        "min_words": 80,
        "placeholder": (
            "Ana Akış:\n"
            "1. Kullanıcı ...\n"
            "2. Sistem ...\n"
            "3. ...\n\n"
            "Alternatif Akış (hata):\n"
            "1a. ...\n"
            "2a. ..."
        ),
    },
    {
        "key": "veri_modeli",
        "title": "10. Veri Modeli",
        "guidance": (
            "Yeni tablo(lar) veya mevcut tabloya eklenecek kolonlar. Her alan için "
            "tipi, boyutu, null durumu, FK ilişkileri yazılır."
        ),
        "required": True,
        "min_words": 60,
        "placeholder": (
            "Yeni Tablo: TabloAdi\n"
            "- Id (INT, PK, IDENTITY)\n"
            "- Alan1 (NVARCHAR(100), NOT NULL)\n"
            "- ...\n\n"
            "İlişkiler:\n"
            "- TabloAdi.MusteriId → Customers.Id (FK)"
        ),
    },
    {
        "key": "arayuz_tasarimi",
        "title": "11. Kullanıcı Arayüzü / Mock-up",
        "guidance": (
            "Hangi ekranlar etkilenecek, hangi yeni ekran/komponentler eklenecek? "
            "Mock-up görseli varsa referans verin, yoksa metin olarak tarif edin."
        ),
        "required": False,
        "min_words": 40,
        "placeholder": (
            "Etkilenen Ekranlar:\n"
            "- Anasayfa (ürün kartı): ...\n"
            "- Ürün detay: ...\n\n"
            "Yeni Ekranlar:\n"
            "- ..."
        ),
    },
    {
        "key": "entegrasyonlar",
        "title": "12. Entegrasyonlar",
        "guidance": (
            "Dış servisler veya iç modüllerle entegrasyon var mı? Email, SMS, ödeme, "
            "kargo, raporlama vs."
        ),
        "required": False,
        "min_words": 30,
        "placeholder": (
            "Email Servisi: ... (varsa servis adı)\n"
            "Ödeme Gateway: ...\n"
            "Yok ise: 'Bu geliştirmede dış entegrasyon yoktur' yazın."
        ),
    },
    {
        "key": "kabul_kriterleri",
        "title": "13. Kabul Kriterleri",
        "guidance": (
            "Bu geliştirme NE olursa 'tamamlandı' sayılır? Her madde test edilebilir, "
            "ölçülebilir olmalı. KK-01, KK-02 şeklinde kodlayın."
        ),
        "required": True,
        "min_words": 60,
        "placeholder": (
            "KK-01: ...\n"
            "KK-02: ...\n"
            "KK-03: ...\n"
            "KK-04: ..."
        ),
    },
    {
        "key": "test_senaryolari",
        "title": "14. Test Senaryoları",
        "guidance": (
            "Pozitif ve negatif senaryolar. Her senaryo: Ön Koşul, Adımlar, Beklenen Sonuç. "
            "TS-01, TS-02 şeklinde kodlayın."
        ),
        "required": True,
        "min_words": 100,
        "placeholder": (
            "TS-01 — [Senaryo Adı]\n"
            "Ön Koşul: ...\n"
            "Adımlar:\n"
            "  1. ...\n"
            "  2. ...\n"
            "Beklenen Sonuç: ...\n\n"
            "TS-02 — ..."
        ),
    },
    {
        "key": "riskler_varsayimlar",
        "title": "15. Riskler ve Varsayımlar",
        "guidance": (
            "Bu projenin başarısını riske atabilecek durumlar ve doğru kabul ettiğimiz "
            "varsayımlar (örn: 'günlük 1000 siparişi aşmayacağız varsayıyoruz')."
        ),
        "required": False,
        "min_words": 30,
        "placeholder": (
            "Riskler:\n"
            "- R-01: ...\n\n"
            "Varsayımlar:\n"
            "- V-01: ..."
        ),
    },
    {
        "key": "acik_sorular",
        "title": "16. Açık Sorular",
        "guidance": (
            "Cevabı henüz bilinmeyen, paydaşlardan netleştirilmesi gereken konular. "
            "Boş bırakmayın; analiz sırasında mutlaka soru çıkar."
        ),
        "required": False,
        "min_words": 20,
        "placeholder": (
            "AS-01: ...?\n"
            "AS-02: ...?"
        ),
    },
]


# Doküman üst kısmındaki sabit alanlar (kapak)
COVER_PLACEHOLDER_REQUEST_TITLE = "[Talep Başlığı]"
COVER_PLACEHOLDER_REQUEST_DESC = "[Talep Açıklaması — buraya rastgele atanan talebin metni gelir]"


def build_template_docx(request_title: str = "", request_description: str = "") -> bytes:
    """
    Boş analiz teknik dokümanı şablonunu Word dosyası (.docx) olarak üretir.
    request_title ve request_description verildiyse kapak sayfasına yazılır.
    """
    doc = Document()

    # --- Sayfa marjları ---
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.2)
        section.right_margin = Cm(2.2)

    # --- Kapak / Başlık ---
    title = doc.add_heading("Analiz Teknik Dokümanı", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = sub.add_run("ErenShop Platformu — İş Analizi Şablonu")
    sub_run.italic = True
    sub_run.font.size = Pt(12)

    doc.add_paragraph()

    # --- Talep özeti kutusu ---
    p_label = doc.add_paragraph()
    r = p_label.add_run("Geliştirme Talebi")
    r.bold = True
    r.font.size = Pt(13)

    p_title = doc.add_paragraph()
    r2 = p_title.add_run("Başlık: ")
    r2.bold = True
    p_title.add_run(request_title or COVER_PLACEHOLDER_REQUEST_TITLE)

    p_desc = doc.add_paragraph()
    r3 = p_desc.add_run("Açıklama: ")
    r3.bold = True
    p_desc.add_run(request_description or COVER_PLACEHOLDER_REQUEST_DESC)

    p_date = doc.add_paragraph()
    r4 = p_date.add_run("Şablon Tarihi: ")
    r4.bold = True
    p_date.add_run(datetime.now().strftime("%d.%m.%Y"))

    doc.add_paragraph()

    # --- Kullanım Notu ---
    note = doc.add_paragraph()
    nr = note.add_run(
        "Not: Her bölümün altındaki italik yazılar yol göstericidir, silebilirsiniz. "
        "Köşeli parantez içindeki [örnek metinleri] kendi içeriğinizle değiştirin. "
        "Bölüm başlıklarını (numara ve isim) DEĞİŞTİRMEYİN — değerlendirme sistemi "
        "başlıkları arıyor."
    )
    nr.italic = True
    nr.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    doc.add_page_break()

    # --- İçindekiler başlığı (otomatik TOC olmadan, basit liste) ---
    doc.add_heading("İçindekiler", level=1)
    for s in TEMPLATE_SECTIONS:
        doc.add_paragraph(s["title"], style="List Number" if False else None)
    doc.add_page_break()

    # --- Bölümler ---
    for section in TEMPLATE_SECTIONS:
        h = doc.add_heading(section["title"], level=1)

        # Guidance (italik gri)
        guidance_p = doc.add_paragraph()
        gr = guidance_p.add_run(section["guidance"])
        gr.italic = True
        gr.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
        gr.font.size = Pt(10)

        if not section.get("required", True):
            opt = doc.add_paragraph()
            opt_r = opt.add_run("(Bu bölüm opsiyoneldir.)")
            opt_r.italic = True
            opt_r.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
            opt_r.font.size = Pt(9)

        # Placeholder içerik
        for line in section["placeholder"].split("\n"):
            doc.add_paragraph(line)

        doc.add_paragraph()  # Boşluk

    # --- Footer (son sayfa) ---
    doc.add_page_break()
    end = doc.add_paragraph()
    end.alignment = WD_ALIGN_PARAGRAPH.CENTER
    er = end.add_run("— Doküman Sonu —")
    er.italic = True
    er.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    # bytes olarak döndür
    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


def get_section_definitions():
    """Evaluation servisi için bölüm tanımlarını döndürür."""
    return TEMPLATE_SECTIONS


# ============================================================
# ÖRNEK (TAM DOLDURULMUŞ) DOKÜMAN
# Eren'in referans alabilmesi için kapsamlı bir örnek talep:
#   "Favori Ürünler Listesi (Wishlist)"
# Her bölüm gerçek bir iş analistinin yazacağı şekilde doldurulmuştur.
# ============================================================
EXAMPLE_REQUEST_TITLE = "Favori Ürünler Listesi (Wishlist)"
EXAMPLE_REQUEST_DESCRIPTION = (
    "Giriş yapan müşteriler beğendikleri ürünleri favorilere ekleyebilmeli. "
    "Profil sayfasında 'Favorilerim' sekmesi olacak. Favorideki ürünün fiyatı "
    "düşerse müşteriye email ile bildirim gitmeli. Favorilerden ürün çıkarılabilmeli. "
    "Bir müşteri en fazla 100 favori ekleyebilir."
)


EXAMPLE_SECTIONS_CONTENT = {
    "dokuman_bilgileri": (
        "Doküman Adı: Favori Ürünler Listesi — Analiz Teknik Dokümanı\n"
        "Versiyon: 1.0\n"
        "Hazırlayan: Eren Sarıteke (İş Analisti)\n"
        "Tarih: 20.05.2026\n"
        "Gözden Geçiren: Ömer Gökdere (Yazılım Yöneticisi)\n"
        "Durum: Onaylı\n"
        "Hedef Sprint: Sprint 28 (03.06.2026 - 17.06.2026)"
    ),

    "yonetici_ozeti": (
        "Bu doküman ErenShop platformuna 'Favori Ürünler Listesi' (Wishlist) özelliğinin "
        "eklenmesini açıklamaktadır. Mevcut sistemde müşteriler beğendikleri ürünleri "
        "kaydedemediği için tekrar bulmakta zorlanmakta, bu durum dönüşüm oranını %12 "
        "düşürmektedir. Önerilen çözüm ile müşteriler ürünleri favorilerine ekleyebilecek, "
        "fiyat düştüğünde otomatik bildirim alacaktır. Bu özelliğin müşteri etkileşimini "
        "%20, tekrar satın alma oranını ise %8 artırması beklenmektedir. Geliştirme süresi "
        "yaklaşık 2 sprint, devreye alma hedefi 30.06.2026'dır."
    ),

    "amac_kapsam": (
        "3.1 Amaç:\n"
        "Müşterilerin ileride satın almayı düşündükleri ürünleri kalıcı olarak "
        "kaydedebilmelerini sağlamak, fiyat düşüşlerinde haberdar ederek satın alma "
        "kararlarını hızlandırmak.\n\n"
        "3.2 Kapsam Dahilinde:\n"
        "- Favori ürün ekleme / çıkarma işlevi (ürün sayfası ve liste sayfasından)\n"
        "- Profil altında 'Favorilerim' sekmesi ve listeleme ekranı\n"
        "- Müşteri başına maksimum 100 favori sınırı\n"
        "- Favorideki ürünün fiyatı düştüğünde email bildirimi (günde 1 mail toplu olarak)\n"
        "- Favori sayısının ürün kartında küçük rozet olarak gösterilmesi (sosyal kanıt)\n"
        "- Favoriden sepete tek tıkla ekleme\n\n"
        "3.3 Kapsam Dışı:\n"
        "- Favori listesini başkasıyla paylaşma (gelecek sprint)\n"
        "- Stok bittiğinde bildirim (ayrı bir talep olarak ele alınacak)\n"
        "- Web push / mobil push bildirimi (sadece email)\n"
        "- Üyeliksiz (anonim) ziyaretçilerin favori eklemesi"
    ),

    "mevcut_durum": (
        "Mevcut ErenShop platformunda müşteriler ürünleri sadece sepete ekleyebilmekte, "
        "ileride satın almak için saklayamamaktadır. Tarayıcı kapatıldığında veya farklı "
        "cihazdan girildiğinde önceden bakılan ürünler kayboluyor. Müşteri analitik "
        "raporlarına göre kullanıcıların %34'ü ilk ziyaretinde alışveriş tamamlamıyor ve "
        "bu kullanıcıların %58'i bir daha aynı ürünü bulamadığı için satın alma fırsatı "
        "kaçıyor. Ayrıca müşteri hizmetlerine en sık gelen sorulardan biri 'geçen sefer "
        "baktığım ürünü bulamıyorum' şikayetidir. Etkilenecek modüller: Müşteri, Ürün, "
        "Bildirim (yeni modül). Mevcut Customers ve Products tablolarına dokunulmayacak, "
        "yeni Favorites tablosu eklenecektir."
    ),

    "paydaslar": (
        "İş Tarafı:\n"
        "- Ürün Sahibi (Product Owner): Ahmet Yıldız (Pazarlama Müdürü)\n"
        "- Sponsor: Genel Müdür Yardımcısı (CRO hedeflerinden sorumlu)\n"
        "- Son Kullanıcı: Üye müşteriler (~12.000 aktif kullanıcı)\n"
        "- Müşteri Hizmetleri: özellik yayına çıkınca destek almak için bilgilendirilecek\n\n"
        "Teknik Taraf:\n"
        "- İş Analisti: Eren Sarıteke\n"
        "- Backend Geliştirici: Burak Kara\n"
        "- Frontend Geliştirici: Selin Aydın\n"
        "- Test Mühendisi: Mert Demir\n"
        "- DevOps: Furkan Şen (email servisi entegrasyonu)\n"
        "- Veritabanı Yöneticisi: Cem Polat (yeni tablo ve indeks onayı)"
    ),

    "fonksiyonel_gereksinimler": (
        "FG-01: Sistem, giriş yapmış müşterinin ürün detay sayfasında 'Favorilere Ekle' "
        "butonu göstermelidir.\n"
        "FG-02: Sistem, favoriye eklenen ürünün butonunu 'Favorilerden Çıkar' olarak "
        "değiştirmelidir.\n"
        "FG-03: Sistem, profil sayfasında 'Favorilerim' sekmesi sunmalı; bu sekmede "
        "müşterinin tüm favori ürünleri ürün kartı formatında listelenmelidir.\n"
        "FG-04: Sistem, favori listesinde her ürün için 'Sepete Ekle' ve 'Favorilerden Çıkar' "
        "aksiyonlarını desteklemelidir.\n"
        "FG-05: Sistem, müşteri başına en fazla 100 favori ürünü saklamalıdır; bu sayıya "
        "ulaşıldığında 'Favori limitiniz dolu' uyarısı vermelidir.\n"
        "FG-06: Sistem, her gece 23:00'te o gün fiyatı düşen favori ürünleri tespit etmeli ve "
        "ilgili müşterilere tek bir özet email göndermelidir.\n"
        "FG-07: Sistem, üyeliksiz ziyaretçilerin favori eklemeye çalışmasında giriş "
        "ekranına yönlendirmelidir (yönlendirme sonrası kullanıcı geri gelirse ürün otomatik "
        "favoriye eklenmelidir).\n"
        "FG-08: Sistem, silinen (IsActive=0) ürünleri favori listesinde 'Artık satışta değil' "
        "etiketi ile göstermeli, üzerine tıklanamamalıdır."
    ),

    "fonksiyonel_olmayan_gereksinimler": (
        "Performans:\n"
        "- Favorilerim sayfası 100 ürünle bile < 2 saniyede yüklenmelidir\n"
        "- 'Favorilere Ekle' API çağrısı < 300 ms cevap vermelidir\n\n"
        "Güvenlik:\n"
        "- Başka müşterinin favori listesi görüntülenememelidir (yetkilendirme)\n"
        "- Favori ekleme isteklerinde rate limit: dakikada 30 ekleme/silme\n\n"
        "Kullanılabilirlik:\n"
        "- Favori ekleme aksiyonu sayfa yenilenmeden gerçekleşmelidir (AJAX)\n"
        "- İşlem başarılı/başarısız olduğunda toast bildirimi gösterilmelidir\n\n"
        "Loglama:\n"
        "- Tüm favori ekleme/çıkarma aksiyonları log tablosuna yazılmalıdır (analitik için)\n\n"
        "Uyumluluk:\n"
        "- KVKK kapsamında müşteri hesabını silerse favori listesi de silinmelidir"
    ),

    "is_kurallari": (
        "IK-01: Bir müşteri aynı ürünü birden fazla kez favorilere ekleyemez.\n"
        "IK-02: Bir müşteri en fazla 100 favori ürün ekleyebilir; bu sınıra ulaşıldığında "
        "yeni ekleme reddedilir.\n"
        "IK-03: Sadece IsActive=1 olan ürünler favoriye eklenebilir.\n"
        "IK-04: Sadece kayıtlı (login olmuş) müşteriler favori ekleyebilir.\n"
        "IK-05: Favori listesi müşteriye özeldir, başka müşteri görüntüleyemez.\n"
        "IK-06: Bir ürün fiyat düştüğünde bildirim sadece SON 24 saatte fiyat değişikliği "
        "olan favori sahiplerine gönderilir; aynı müşteriye aynı ürün için günde 1 mail.\n"
        "IK-07: Fiyat düşüşü en az %5 olmalı; bunun altındaki değişiklikler bildirim "
        "tetiklemez (gürültüyü önlemek için).\n"
        "IK-08: Müşteri hesabını sildiğinde tüm favorileri de silinir (CASCADE)."
    ),

    "surec_akisi": (
        "Ana Akış — Favoriye Ürün Ekleme:\n"
        "1. Müşteri ürün detay sayfasını açar\n"
        "2. 'Favorilere Ekle' butonuna tıklar\n"
        "3. Sistem müşterinin giriş yapıp yapmadığını kontrol eder\n"
        "4. Müşteri giriş yapmışsa: Favori sayısı 100'ün altındaysa Favorites tablosuna kayıt eklenir\n"
        "5. Buton 'Favorilerden Çıkar' olarak değişir, toast: 'Favorilere eklendi ❤️'\n\n"
        "Alternatif Akış A — Giriş yapılmamış:\n"
        "3a. Sistem müşteriyi giriş ekranına yönlendirir, geri dönüş URL'sini saklar\n"
        "3b. Giriş başarılı olunca otomatik favoriye eklenir ve ürün sayfasına döner\n\n"
        "Alternatif Akış B — Limit dolu:\n"
        "4a. Favori sayısı 100 ise sistem hata mesajı gösterir: 'Favori limitiniz dolu (100/100)'\n"
        "4b. 'Favorilerden çıkar' linki ile profil sayfasına yönlendirme önerilir\n\n"
        "Fiyat Düşüşü Email Akışı (Zamanlanmış İş - her gece 23:00):\n"
        "1. Sistem son 24 saatte fiyatı en az %5 düşen ürünleri bulur\n"
        "2. Bu ürünleri favoriye eklemiş tüm müşterileri tespit eder\n"
        "3. Her müşteri için tek bir özet email oluşturur (birden fazla ürün varsa hepsi tek mailde)\n"
        "4. Email gönderim servisine kuyruğa atar\n"
        "5. Gönderim sonucu Notifications log tablosuna yazılır"
    ),

    "veri_modeli": (
        "Yeni Tablo: Favorites\n"
        "- Id (INT, PK, IDENTITY)\n"
        "- CustomerId (INT, NOT NULL, FK → Customers.Id)\n"
        "- ProductId (INT, NOT NULL, FK → Products.Id)\n"
        "- CreatedAt (DATETIME, NOT NULL, DEFAULT GETDATE())\n"
        "- LastNotifiedAt (DATETIME, NULL) — son fiyat düşüş bildiriminin zamanı (aynı gün tekrar yollanmasın diye)\n"
        "- UNIQUE Constraint: (CustomerId, ProductId) — IK-01 garantisi\n"
        "- INDEX: (CustomerId) — Favorilerim sorgusu için\n"
        "- INDEX: (ProductId) — fiyat düşüşünde kim haberdar olacak sorgusu için\n\n"
        "Mevcut Tabloya Eklenecek Kolon:\n"
        "- Products.PriceLastChangedAt (DATETIME, NULL) — fiyat değişim zamanını izlemek için\n\n"
        "İlişkiler:\n"
        "- Favorites.CustomerId → Customers.Id (ON DELETE CASCADE — IK-08)\n"
        "- Favorites.ProductId → Products.Id (NO CASCADE — ürün silinmesin, sadece deaktive)\n\n"
        "Yeni Log Tablosu (analitik için):\n"
        "Yok — Notifications tablosu var, oraya yazılabilir."
    ),

    "arayuz_tasarimi": (
        "Etkilenen Ekranlar:\n"
        "- Ürün Detay Sayfası: 'Sepete Ekle' butonunun yanına kalp ikonu eklenecek. Doluysa "
        "favoride, boşsa değil. Tıklayınca animasyonlu doldurma/boşaltma.\n"
        "- Ürün Kartı (liste/grid): Sağ üst köşede küçük kalp ikonu. Mobilde de tıklanabilir.\n"
        "- Profil Sayfası: Üst menüye 'Favorilerim (12)' sekmesi eklenecek. Parantez içinde sayı.\n"
        "- Header: Sepet ikonunun yanına küçük kalp+sayı rozeti.\n\n"
        "Yeni Ekranlar:\n"
        "- Favorilerim Listesi: Ürün kartlarını grid olarak gösterir, her kartta 'Sepete Ekle' "
        "ve 'Çıkar' butonları. Boşsa ilüstrasyon + 'Henüz favori eklemediniz' mesajı + "
        "'Ürünleri Keşfet' butonu.\n\n"
        "Mock-up linki: Figma → ErenShop / Wishlist v1.2 (Selin Aydın hazırlayacak)\n\n"
        "Email Şablonu:\n"
        "- Konu: 'Beğendiğin ürünlerde indirim var! 💚'\n"
        "- Gövde: 'Merhaba {ad}, favori listendeki şu ürünlerin fiyatı düştü:' + ürün listesi "
        "+ eski/yeni fiyat + 'Şimdi Al' butonu"
    ),

    "entegrasyonlar": (
        "Email Servisi:\n"
        "- Mevcut Notifications modülü kullanılacak (zaten sipariş onay maillerini bu modül gönderiyor).\n"
        "- Yeni şablon ID: WISHLIST_PRICE_DROP\n\n"
        "Zamanlanmış İş (Scheduler):\n"
        "- Mevcut cron servisine yeni iş eklenecek: gece 23:00 → price-drop-notifier.py\n"
        "- DevOps (Furkan) cron tanımını ekleyecek\n\n"
        "Analytics:\n"
        "- Google Analytics custom event: 'add_to_wishlist' ve 'remove_from_wishlist'\n"
        "- Frontend (Selin) tarafından gönderilecek\n\n"
        "Üçüncü Parti Bağımlılık:\n"
        "- Yok. Mevcut altyapıyla çalışır."
    ),

    "kabul_kriterleri": (
        "KK-01: Müşteri giriş yapmışken bir ürünün favori ikonuna tıkladığında ürün anında "
        "favorilere eklenmeli (< 300 ms cevap), buton görseli değişmeli.\n"
        "KK-02: Aynı müşteri aynı ürünü iki kez favoriye eklemeye çalıştığında ikinci istek "
        "hata vermeli (UNIQUE constraint'ten gelen 400 Bad Request).\n"
        "KK-03: Profil > Favorilerim sayfası 100 ürünle birlikte 2 saniyenin altında yüklenmeli.\n"
        "KK-04: Favori sayısı 100'e ulaşmış müşteri yeni ürün eklemeye çalıştığında 'Limit dolu' "
        "mesajı görmeli, kayıt eklenmemiş olmalı.\n"
        "KK-05: Bir ürünün fiyatı %6 düşürüldüğünde, ertesi gün 23:00'te o ürünü favorisine "
        "eklemiş tüm aktif müşterilere SADECE BİR mail gitmeli.\n"
        "KK-06: Bir ürünün fiyatı %3 düşürüldüğünde HİÇ mail gitmemeli (IK-07).\n"
        "KK-07: Müşteri hesabı silindiğinde tüm favori kayıtları silinmeli (Favorites tablosu boşalır).\n"
        "KK-08: Üyeliksiz ziyaretçi favori eklemeye çalıştığında giriş sayfasına yönlendirilmeli "
        "ve giriş sonrası otomatik favoriye eklenmeli."
    ),

    "test_senaryolari": (
        "TS-01 — Standart favoriye ekleme (Happy Path)\n"
        "Ön Koşul: Müşteri ID=5 sisteme giriş yapmış. Ürün ID=12 aktif ve stoğu var. Müşterinin favorilerinde Id=12 YOK.\n"
        "Adımlar:\n"
        "  1. Ürün detay sayfasını aç (/products/12)\n"
        "  2. Favori (kalp) ikonuna tıkla\n"
        "Beklenen Sonuç: Toast 'Favorilere eklendi' görünür. Veritabanında Favorites tablosuna (CustomerId=5, ProductId=12) eklenir. Kalp ikonu dolu görünür.\n\n"
        "TS-02 — Aynı ürünü tekrar ekleme (negatif)\n"
        "Ön Koşul: Müşteri 5'in favorilerinde Ürün 12 var.\n"
        "Adımlar:\n"
        "  1. Postman'dan POST /api/favorites { customerId: 5, productId: 12 }\n"
        "Beklenen Sonuç: 400 Bad Request, message: 'Bu ürün zaten favorilerinizde.'\n\n"
        "TS-03 — Limit kontrolü\n"
        "Ön Koşul: Müşteri 5'in 100 favori ürünü var.\n"
        "Adımlar:\n"
        "  1. 101. ürünü favoriye eklemeye çalış\n"
        "Beklenen Sonuç: 400 Bad Request, message: 'Favori limitiniz dolu (100/100)'.\n\n"
        "TS-04 — Pasif (silinmiş) ürün ekleme (negatif)\n"
        "Ön Koşul: Ürün 99 IsActive=0.\n"
        "Adımlar:\n"
        "  1. POST /api/favorites { customerId: 5, productId: 99 }\n"
        "Beklenen Sonuç: 400 Bad Request, message: 'Bu ürün artık satışta değil.'\n\n"
        "TS-05 — Fiyat düşüş emaili (Happy Path)\n"
        "Ön Koşul: Ürün 12'nin fiyatı 1000 TL. Müşteri 5'in favorilerinde Ürün 12 var.\n"
        "Adımlar:\n"
        "  1. Ürün 12'nin fiyatını 900 TL yap (%10 düşüş)\n"
        "  2. Gece 23:00 cron'unu manuel tetikle\n"
        "Beklenen Sonuç: Müşteri 5'e tek email gider, içeriğinde Ürün 12'nin eski (1000) ve yeni (900) fiyatı yer alır. Favorites.LastNotifiedAt güncellenir.\n\n"
        "TS-06 — Fiyat düşüşü %5 altı (negatif)\n"
        "Ön Koşul: Ürün 12 = 1000 TL, Müşteri 5'in favorisinde.\n"
        "Adımlar:\n"
        "  1. Fiyatı 970 TL yap (%3 düşüş)\n"
        "  2. Cron'u tetikle\n"
        "Beklenen Sonuç: HİÇ email gönderilmez. Notifications tablosunda yeni kayıt YOK.\n\n"
        "TS-07 — Hesap silindiğinde CASCADE\n"
        "Ön Koşul: Müşteri 5'in 25 favori ürünü var.\n"
        "Adımlar:\n"
        "  1. Müşteri 5'i sil\n"
        "Beklenen Sonuç: Favorites tablosunda Müşteri 5'e ait hiç kayıt kalmaz."
    ),

    "riskler_varsayimlar": (
        "Riskler:\n"
        "- R-01: Cron iş başarısız olursa müşteriler bildirim alamaz. Çözüm: monitoring + Slack alarmı.\n"
        "- R-02: Yüksek trafikte (kampanya günü) email kuyruk birikebilir. Çözüm: kuyruk metriklerini izle, gerekirse worker sayısını artır.\n"
        "- R-03: Favori tablosu büyüdüğünde sorgu yavaşlayabilir. Çözüm: tanımlanan indeksler ve 6 ayda bir bakım planı.\n\n"
        "Varsayımlar:\n"
        "- V-01: Email servisi (mevcut SMTP) günde 50.000 mail kapasiteyi karşılayacak.\n"
        "- V-02: Aktif kullanıcıların ortalama 8 favori ekleyeceği tahmin ediliyor (toplam ~96K kayıt).\n"
        "- V-03: Üyelik gerektiren bir özellik olduğu için yeni üye kazanımına katkı yapacak (henüz ölçülmedi)."
    ),

    "acik_sorular": (
        "AS-01: Üyeliksiz ziyaretçi 'favorilere ekle' tıkladığında giriş yerine kayıt da önerilsin mi?\n"
        "AS-02: Favori limiti 100 yeterli mi, B2B müşteriler için daha yüksek olmalı mı? Pazarlamaya soruldu.\n"
        "AS-03: Fiyat düşüş eşiği %5 sabit mi, müşteri profilinde özelleştirilebilir mi olsun?\n"
        "AS-04: Mobil app de aynı özelliği desteklemesi gerekecek — mobil ekibi paralel mi gidecek yoksa sonra mı?\n"
        "AS-05: Stok bittiğinde favoride duran ürün için 'tekrar gelince haber ver' otomatik mi aktif olsun? (Ayrı bir talep, ama bağımlılık var)"
    ),
}


def build_example_docx() -> bytes:
    """
    Tam doldurulmuş örnek analiz dokümanı (Favori Listesi talebi için).
    Eren bunu indirip referans olarak kullanır, kendi talebi için ayrı bir
    boş şablon doldurur.
    """
    doc = Document()

    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.2)
        section.right_margin = Cm(2.2)

    # --- Kapak ---
    title = doc.add_heading("Analiz Teknik Dokümanı", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = sub.add_run("ÖRNEK / REFERANS DOKÜMAN — Tam Doldurulmuş")
    sr.bold = True
    sr.italic = True
    sr.font.size = Pt(12)
    sr.font.color.rgb = RGBColor(0x6c, 0x63, 0xff)

    doc.add_paragraph()

    # --- Bilgilendirme kutusu ---
    info = doc.add_paragraph()
    ir = info.add_run(
        "⚠️ Bu doküman BİR ÖRNEKTİR. Aşağıdaki 'Favori Ürünler Listesi' talebi için "
        "tam doldurulmuş şekilde hazırlanmıştır. Sen kendi sana atanan talebin için "
        "BOŞ ŞABLONU indir ve onu doldur. Bu örneği rehber olarak kullan: bölüm "
        "uzunlukları, detay seviyesi ve yazım stili için referans alabilirsin."
    )
    ir.italic = True
    ir.font.color.rgb = RGBColor(0xc6, 0x28, 0x28)
    ir.font.size = Pt(10)

    doc.add_paragraph()

    # --- Talep kutusu ---
    p_label = doc.add_paragraph()
    r = p_label.add_run("Geliştirme Talebi (Örnek)")
    r.bold = True
    r.font.size = Pt(13)

    p_title = doc.add_paragraph()
    r2 = p_title.add_run("Başlık: ")
    r2.bold = True
    p_title.add_run(EXAMPLE_REQUEST_TITLE)

    p_desc = doc.add_paragraph()
    r3 = p_desc.add_run("Açıklama: ")
    r3.bold = True
    p_desc.add_run(EXAMPLE_REQUEST_DESCRIPTION)

    p_date = doc.add_paragraph()
    r4 = p_date.add_run("Doküman Tarihi: ")
    r4.bold = True
    p_date.add_run(datetime.now().strftime("%d.%m.%Y"))

    doc.add_page_break()

    # --- İçindekiler ---
    doc.add_heading("İçindekiler", level=1)
    for s in TEMPLATE_SECTIONS:
        doc.add_paragraph(s["title"])
    doc.add_page_break()

    # --- Bölümler (her birinde tam doldurulmuş örnek içerik) ---
    for section in TEMPLATE_SECTIONS:
        doc.add_heading(section["title"], level=1)

        # Tam doldurulmuş içerik
        content = EXAMPLE_SECTIONS_CONTENT.get(section["key"], "")
        if content:
            for line in content.split("\n"):
                doc.add_paragraph(line)
        else:
            # Fallback (olmamalı)
            doc.add_paragraph("(Bu bölüm için örnek içerik henüz yazılmadı.)")

        doc.add_paragraph()

    # --- Footer ---
    doc.add_page_break()
    end = doc.add_paragraph()
    end.alignment = WD_ALIGN_PARAGRAPH.CENTER
    er = end.add_run("— Örnek Doküman Sonu —")
    er.italic = True
    er.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()
